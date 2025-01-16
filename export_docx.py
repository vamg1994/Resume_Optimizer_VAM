from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_resume_docx(structured_cv, output_path='resume.docx', config=None):
    """Generate DOCX resume from structured CV data with custom configurations"""
    try:
        doc = Document()
        
        # Default configuration
        default_config = {
            'name_size': 24,
            'contact_size': 10,
            'heading_size': 14,
            'title_size': 12,
            'body_size': 11,
            'margins': 0.5,
            'line_spacing': 1.15
        }
        
        # Use provided config or defaults
        config = config or default_config
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(config['margins'])
            section.bottom_margin = Inches(config['margins'])
            section.left_margin = Inches(config['margins'])
            section.right_margin = Inches(config['margins'])

        # Add Name
        name = doc.add_paragraph()
        name_run = name.add_run(structured_cv['name'])
        name_run.bold = True
        name_run.font.size = Pt(config['name_size'])
        name.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add Contact Info
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact.add_run(' | '.join(structured_cv['contact'])).font.size = Pt(config['contact_size'])
        
        # Set line spacing for the document
        for paragraph in doc.paragraphs:
            paragraph.paragraph_format.line_spacing = config['line_spacing']

        # Add Professional Summary
        doc.add_heading('Professional Summary', level=1)
        summary = doc.add_paragraph()
        summary.add_run(structured_cv['professional_summary']).font.size = Pt(config['body_size'])

        # Add Work Experience
        doc.add_heading('Work Experience', level=1)
        for job in structured_cv['work_experience']:
            # Job title
            p = doc.add_paragraph()
            title_run = p.add_run(job['title'])
            title_run.bold = True
            title_run.font.size = Pt(config['title_size'])
            
            # Company and dates
            company = doc.add_paragraph()
            company.add_run(f"{job['company']} | {job['dates']}").font.size = Pt(config['body_size'])
            
            # Responsibilities
            for resp in job['responsibilities']:
                bullet = doc.add_paragraph(style='List Bullet')
                bullet.add_run(resp).font.size = Pt(config['body_size'])

        # Add Education
        doc.add_heading('Education', level=1)
        for edu in structured_cv['education']:
            # Degree
            p = doc.add_paragraph()
            degree_run = p.add_run(edu['degree'])
            degree_run.bold = True
            degree_run.font.size = Pt(config['title_size'])
            
            # Institution and dates
            inst = doc.add_paragraph()
            inst.add_run(f"{edu['institution']} | {edu['dates']}").font.size = Pt(config['body_size'])
            
            # Details
            for detail in edu['details']:
                bullet = doc.add_paragraph(style='List Bullet')
                bullet.add_run(detail).font.size = Pt(config['body_size'])

        # Add Skills
        doc.add_heading('Skills', level=1)
        for category, skills in structured_cv['skills'].items():
            # Category
            p = doc.add_paragraph()
            category_run = p.add_run(category)
            category_run.bold = True
            category_run.font.size = Pt(config['title_size'])
            
            # Skills list
            skills_p = doc.add_paragraph()
            skills_p.add_run(', '.join(skills)).font.size = Pt(config['body_size'])

        # Save the document
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Error generating DOCX: {str(e)}")
        return False 